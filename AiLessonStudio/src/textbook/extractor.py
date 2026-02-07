import re
import json
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
import pdfplumber
import fitz
from docx import Document


@dataclass
class ExtractedContent:
    """Structured content extraction result"""
    text: str
    chapters: Dict[str, Dict[str, Any]]
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    errors: List[str]


class ContentExtractor:
    """Advanced content extractor for various file formats"""

    def __init__(self, config):
        self.config = config

        # Initialize extractors
        self._init_extractors()
        print("✅ Content Extractor initialized")

    def _init_extractors(self):
        """Initialize format-specific extractors"""
        self.extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.txt': self._extract_txt,
            '.json': self._extract_json
        }

    def extract_content(self, file_path: str,
                        extract_images: bool = False,
                        extract_tables: bool = True) -> ExtractedContent:
        """Extract content from file"""
        import os
        from pathlib import Path

        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension not in self.extractors:
            return ExtractedContent(
                text="",
                chapters={},
                metadata={'error': f'Unsupported file format: {extension}'},
                tables=[],
                images=[],
                errors=[f'Unsupported file format: {extension}']
            )

        try:
            # Call appropriate extractor
            extractor = self.extractors[extension]
            result = extractor(file_path, extract_images, extract_tables)

            # Post-process extracted content
            result = self._post_process_content(result)

            return result

        except Exception as e:
            return ExtractedContent(
                text="",
                chapters={},
                metadata={'error': str(e)},
                tables=[],
                images=[],
                errors=[f'Extraction error: {str(e)}']
            )

    def _extract_pdf(self, file_path: Path,
                     extract_images: bool,
                     extract_tables: bool) -> ExtractedContent:
        """Extract content from PDF"""
        text_parts = []
        chapters = {}
        tables = []
        images = []
        metadata = {}

        try:
            # Method 1: Use PyMuPDF for text and metadata
            with fitz.open(file_path) as doc:
                metadata = {
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'pages': len(doc)
                }

                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

                    if extract_images:
                        image_list = page.get_images()
                        for img in image_list:
                            images.append({
                                'page': page_num + 1,
                                'index': len(images),
                                'xref': img[0]
                            })

            # Method 2: Use pdfplumber for tables and better text
            if extract_tables:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        page_tables = page.extract_tables()
                        for table_num, table in enumerate(page_tables):
                            if table:
                                tables.append({
                                    'page': page_num + 1,
                                    'table_num': table_num,
                                    'data': table,
                                    'rows': len(table),
                                    'columns': len(table[0]) if table[0] else 0
                                })

            # Combine text
            full_text = "\n\n".join(text_parts)

            # Extract chapters
            chapters = self._extract_chapters_from_text(full_text)

            return ExtractedContent(
                text=full_text,
                chapters=chapters,
                metadata=metadata,
                tables=tables,
                images=images,
                errors=[]
            )

        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")

    def _extract_docx(self, file_path: Path,
                      extract_images: bool,
                      extract_tables: bool) -> ExtractedContent:
        """Extract content from DOCX"""
        doc = Document(file_path)

        text_parts = []
        chapters = {}
        tables = []

        # Extract text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        if extract_tables:
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    tables.append({
                        'table_num': table_num,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data[0] else 0
                    })

        full_text = "\n".join(text_parts)

        # Extract chapters
        chapters = self._extract_chapters_from_text(full_text)

        return ExtractedContent(
            text=full_text,
            chapters=chapters,
            metadata={'format': 'docx', 'paragraphs': len(doc.paragraphs)},
            tables=tables,
            images=[],  # DOCX image extraction is complex
            errors=[]
        )

    def _extract_txt(self, file_path: Path,
                     extract_images: bool,
                     extract_tables: bool) -> ExtractedContent:
        """Extract content from TXT"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            full_text = f.read()

        # Extract chapters
        chapters = self._extract_chapters_from_text(full_text)

        return ExtractedContent(
            text=full_text,
            chapters=chapters,
            metadata={'format': 'txt', 'size': len(full_text)},
            tables=[],
            images=[],
            errors=[]
        )

    def _extract_json(self, file_path: Path,
                      extract_images: bool,
                      extract_tables: bool) -> ExtractedContent:
        """Extract content from JSON"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert JSON to text
        text_parts = []

        def process_value(key, value, indent=0):
            indent_str = "  " * indent
            if isinstance(value, dict):
                text_parts.append(f"{indent_str}{key}:")
                for k, v in value.items():
                    process_value(k, v, indent + 1)
            elif isinstance(value, list):
                text_parts.append(f"{indent_str}{key}:")
                for i, item in enumerate(value):
                    text_parts.append(f"{indent_str}  [{i}]: {str(item)}")
            else:
                text_parts.append(f"{indent_str}{key}: {value}")

        for key, value in data.items():
            process_value(key, value)

        full_text = "\n".join(text_parts)

        return ExtractedContent(
            text=full_text,
            chapters={},
            metadata={'format': 'json', 'keys': list(data.keys())},
            tables=[],
            images=[],
            errors=[]
        )

    def _extract_chapters_from_text(self, text: str) -> Dict[str, Dict[str, Any]]:
        """Extract chapters from text using multiple strategies"""
        chapters = {}

        # Strategy 1: Numbered chapters (Chapter 1, Chapter 2, etc.)
        pattern1 = r'(?i)(?:^|\n)\s*(?:chapter|chap|ch\.?)\s+(\d+(?:\.\d+)?)[:.\s]+\s*(.+?)(?=\n\s*(?:chapter|chap|ch\.?)\s+\d+|\Z)'
        matches1 = list(re.finditer(pattern1, text, re.DOTALL))

        # Strategy 2: Numbered sections (1. Introduction, 2. Background, etc.)
        pattern2 = r'(?i)(?:^|\n)\s*(\d+(?:\.\d+)*)\.\s+(.+?)(?=\n\s*\d+(?:\.\d+)*\.\s+|\Z)'
        matches2 = list(re.finditer(pattern2, text, re.DOTALL))

        # Strategy 3: Major headings (all caps or with special formatting)
        pattern3 = r'(?:^|\n)\s*([A-Z][A-Z\s]{10,})(?:\n|$)'
        matches3 = list(re.finditer(pattern3, text))

        # Choose best strategy
        all_matches = []
        if len(matches1) >= 2:
            all_matches = matches1
            match_type = 'chapter'
        elif len(matches2) >= 3:
            all_matches = matches2
            match_type = 'section'
        elif len(matches3) >= 2:
            all_matches = matches3
            match_type = 'heading'
        else:
            # No clear structure, create single chapter
            chapters['chapter_1'] = {
                'title': 'Full Content',
                'number': '1',
                'content': text[:5000],  # Limit content
                'pages': '1-?',
                'type': 'single'
            }
            return chapters

        # Process matches
        for i, match in enumerate(all_matches):
            if match_type == 'chapter':
                chapter_num = match.group(1)
                chapter_title = match.group(2).strip()
                chapter_content = match.group(0)
            elif match_type == 'section':
                chapter_num = match.group(1)
                chapter_title = match.group(2).strip()
                chapter_content = match.group(0)
            else:  # heading
                chapter_num = str(i + 1)
                chapter_title = match.group(1).strip()
                # Get content until next heading
                start = match.end()
                next_match = all_matches[i + 1] if i + 1 < len(all_matches) else None
                end = next_match.start() if next_match else len(text)
                chapter_content = text[start:end].strip()

            chapter_id = f"{match_type}_{chapter_num.replace('.', '_')}"

            chapters[chapter_id] = {
                'title': chapter_title,
                'number': chapter_num,
                'content': chapter_content[:3000],  # Limit content
                'type': match_type,
                'match_score': len(chapter_content) / 1000  # Simple scoring
            }

        return chapters

    def _post_process_content(self, content: ExtractedContent) -> ExtractedContent:
        """Post-process extracted content"""
        # Clean up text
        if content.text:
            # Remove excessive whitespace
            content.text = re.sub(r'\n\s*\n\s*\n', '\n\n', content.text)

            # Remove page numbers and headers
            content.text = re.sub(r'Page\s+\d+\s+', '', content.text)
            content.text = re.sub(r'\d+\s+of\s+\d+', '', content.text)

        # Enhance chapter information
        for chapter_id, chapter in content.chapters.items():
            # Add word count
            if 'content' in chapter:
                word_count = len(chapter['content'].split())
                chapter['word_count'] = word_count

            # Add estimated reading time
            if 'word_count' in chapter:
                reading_time = chapter['word_count'] / 200  # 200 words per minute
                chapter['reading_time_min'] = round(reading_time, 1)

        return content

    def analyze_text_structure(self, text: str) -> Dict[str, Any]:
        """Analyze text structure and quality"""
        if not text:
            return {'error': 'No text provided'}

        # Basic metrics
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        paragraphs = text.split('\n\n')

        metrics = {
            'word_count': len(words),
            'sentence_count': len(sentences),
            'paragraph_count': len(paragraphs),
            'avg_sentence_length': len(words) / max(1, len(sentences)),
            'avg_word_length': sum(len(word) for word in words) / max(1, len(words))
        }

        # Detect structure
        structure = {
            'has_chapters': bool(re.search(r'(?i)chapter\s+\d+', text)),
            'has_sections': bool(re.search(r'\d+\.\s+[A-Z]', text)),
            'has_headings': bool(re.search(r'^[A-Z][A-Z\s]{5,}$', text, re.MULTILINE)),
            'has_lists': bool(re.search(r'(?:^|\n)[•\-\*]\s', text)),
            'has_tables': bool(re.search(r'\+[-]+\+', text))  # Simple table detection
        }

        # Quality indicators
        quality = {
            'readability_score': self._calculate_readability(text),
            'technical_term_count': self._count_technical_terms(text),
            'cloud_term_density': self._calculate_cloud_term_density(text)
        }

        return {
            'metrics': metrics,
            'structure': structure,
            'quality': quality,
            'recommendations': self._generate_structure_recommendations(structure, metrics)
        }

    def _calculate_readability(self, text: str) -> float:
        """Calculate readability score (simplified)"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)

        if len(words) == 0 or len(sentences) == 0:
            return 0.0

        # Simple Flesch Reading Ease approximation
        avg_sentence_length = len(words) / len(sentences)
        avg_word_length = sum(len(word) for word in words) / len(words)

        # Normalize to 0-100 scale
        score = 100 - (avg_sentence_length * 1.5 + avg_word_length * 10)
        return max(0.0, min(100.0, score))

    def _count_technical_terms(self, text: str) -> int:
        """Count technical terms in text"""
        technical_terms = [
            'cloud', 'computing', 'server', 'virtual', 'container',
            'kubernetes', 'docker', 'aws', 'azure', 'google',
            'infrastructure', 'platform', 'software', 'service',
            'scalability', 'availability', 'reliability', 'security',
            'deployment', 'configuration', 'orchestration', 'automation'
        ]

        text_lower = text.lower()
        count = 0
        for term in technical_terms:
            count += text_lower.count(term)

        return count

    def _calculate_cloud_term_density(self, text: str) -> float:
        """Calculate cloud terminology density"""
        words = text.split()
        if not words:
            return 0.0

        cloud_terms = self._count_technical_terms(text)
        return cloud_terms / len(words)

    def _generate_structure_recommendations(self, structure: Dict,
                                            metrics: Dict) -> List[str]:
        """Generate recommendations for improving text structure"""
        recommendations = []

        if not structure['has_chapters'] and metrics['word_count'] > 1000:
            recommendations.append("Consider adding chapter headings for better organization")

        if metrics['avg_sentence_length'] > 25:
            recommendations.append("Try using shorter sentences for better readability")

        if structure['cloud_term_density'] < 0.05:
            recommendations.append("Consider adding more cloud-specific terminology")

        if not structure['has_lists']:
            recommendations.append("Use bullet points or numbered lists for key concepts")

        return recommendations[:3]